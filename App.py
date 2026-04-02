from docx import Document
from flask import Flask, jsonify, send_from_directory, request
import os, re, json, uuid
from werkzeug.utils import secure_filename

app = Flask(__name__)

# ==================== CONFIG ====================
PHOTOS_FOLDER   = "photos"
PHOTOS_META     = "photos/photos.json"
LETTERS_FOLDER  = "letters"

ALLOWED_EXTENSIONS = {"jpg", "jpeg", "png", "gif", "webp", "heic"}
MAX_CONTENT_LENGTH = 20 * 1024 * 1024   # 20 MB per file

app.config["MAX_CONTENT_LENGTH"] = MAX_CONTENT_LENGTH

# Ensure folders exist
os.makedirs(PHOTOS_FOLDER, exist_ok=True)
os.makedirs(LETTERS_FOLDER, exist_ok=True)

def allowed_file(filename):
    return "." in filename and filename.rsplit(".", 1)[1].lower() in ALLOWED_EXTENSIONS

# ==================== LOAD / SAVE PHOTO META ====================
def load_photo_meta():
    if not os.path.exists(PHOTOS_META):
        return []
    with open(PHOTOS_META, "r", encoding="utf-8") as f:
        return json.load(f)

def save_photo_meta(data):
    os.makedirs(os.path.dirname(PHOTOS_META), exist_ok=True)
    with open(PHOTOS_META, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)

# ==================== STATIC FILES ====================
@app.route("/")
def home():
    return send_from_directory(".", "page1.html")

@app.route("/page2.css")
def css():
    return send_from_directory(".", "page2.css")

@app.route("/page3.js")
def js():
    return send_from_directory(".", "page3.js")

# Serve uploaded photos
@app.route("/photos/<filename>")
def serve_photo(filename):
    return send_from_directory(PHOTOS_FOLDER, filename)

# ==================== LETTERS ====================
@app.route("/letters")
def get_letters():
    letters_list = []
    for filename in os.listdir(LETTERS_FOLDER):
        if not filename.endswith(".docx"):
            continue
        file_path = os.path.join(LETTERS_FOLDER, filename)
        doc = Document(file_path)
        full_text = [para.text for para in doc.paragraphs]
        title   = full_text[0] if len(full_text) > 0 else filename
        date    = full_text[1] if len(full_text) > 1 else "Unknown date"
        content = "\n".join(full_text[2:]) if len(full_text) > 2 else "\n".join(full_text)
        letters_list.append({
            "title": title,
            "date": date,
            "content": content,
            "filename": filename
        })

    def extract_number(title):
        match = re.search(r"(\d+)", title)
        return int(match.group(1)) if match else float("inf")

    letters_list.sort(key=lambda x: extract_number(x["title"]))
    return jsonify(letters_list)

@app.route("/rename_letter", methods=["POST"])
def rename_letter():
    data      = request.get_json()
    filename  = data.get("filename")
    new_title = data.get("new_title")
    if not filename or not new_title:
        return jsonify({"success": False, "error": "Invalid data"})
    file_path = os.path.join(LETTERS_FOLDER, filename)
    if not os.path.exists(file_path):
        return jsonify({"success": False, "error": "File not found"})
    doc = Document(file_path)
    if doc.paragraphs:
        doc.paragraphs[0].text = new_title
        doc.save(file_path)
        return jsonify({"success": True})
    return jsonify({"success": False, "error": "No paragraphs found"})

# ==================== PHOTOS — GET ====================
@app.route("/photos")
def get_photos():
    meta = load_photo_meta()
    existing = [
        {**p, "url": f"/photos/{p['filename']}"}
        for p in meta
        if os.path.exists(os.path.join(PHOTOS_FOLDER, p["filename"]))
    ]
    return jsonify(existing)

# ==================== PHOTOS — UPLOAD ====================
@app.route("/upload_photo", methods=["POST"])
def upload_photo():
    if "photo" not in request.files:
        return jsonify({"success": False, "error": "No file provided"})

    file    = request.files["photo"]
    caption = request.form.get("caption", "").strip()

    if file.filename == "":
        return jsonify({"success": False, "error": "Empty filename"})

    if not allowed_file(file.filename):
        return jsonify({"success": False, "error": "File type not allowed"})

    # Unique safe filename
    ext          = file.filename.rsplit(".", 1)[1].lower()
    unique_name  = f"{uuid.uuid4().hex}.{ext}"
    save_path    = os.path.join(PHOTOS_FOLDER, unique_name)
    file.save(save_path)

    # Update metadata
    meta = load_photo_meta()
    meta.append({
        "filename": unique_name,
        "caption":  caption or "A memory 🌸",
        "original": secure_filename(file.filename)
    })
    save_photo_meta(meta)

    return jsonify({"success": True, "filename": unique_name})

# ==================== PHOTOS — DELETE ====================
@app.route("/delete_photo", methods=["POST"])
def delete_photo():
    data     = request.get_json()
    filename = data.get("filename", "")

    safe = secure_filename(filename)
    if not safe or safe != filename:
        return jsonify({"success": False, "error": "Invalid filename"})

    file_path = os.path.join(PHOTOS_FOLDER, safe)
    if os.path.exists(file_path):
        os.remove(file_path)

    meta = load_photo_meta()
    meta = [p for p in meta if p["filename"] != safe]
    save_photo_meta(meta)

    return jsonify({"success": True})

# ==================== RUN ====================
if __name__ == "__main__":
    app.run(debug=True)